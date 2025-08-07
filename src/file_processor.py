"""
File Processing Utilities
Handle document and Excel file processing for SRF extraction
"""
import os
import tempfile
from typing import Dict, List, Optional, Tuple
import pandas as pd
from docx import Document
import logging

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles file processing for SRF content and supporting information"""
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Dict:
        """Extract text content from DOCX file"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Read document
                doc = Document(tmp_file_path)
                
                # Extract text from paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())
                
                # Extract text from tables
                tables_text = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_data.append(cell.text.strip())
                        if row_data:
                            table_data.append(" | ".join(row_data))
                    if table_data:
                        tables_text.append("\n".join(table_data))
                
                # Combine all text
                all_text = "\n\n".join(paragraphs)
                if tables_text:
                    all_text += "\n\nTables:\n" + "\n\n".join(tables_text)
                
                return {
                    'success': True,
                    'text': all_text,
                    'paragraphs_count': len(paragraphs),
                    'tables_count': len(tables_text)
                }
                
            finally:
                # Clean up temp file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                    
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'text': ''
            }
    
    @staticmethod
    def extract_text_from_doc(file_content: bytes) -> Dict:
        """Extract text from DOC file using multiple methods"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                extracted_text = ""
                method_used = ""
                
                # Method 1: Try using win32com (Windows only - most reliable for .doc)
                if os.name == 'nt':
                    try:
                        import win32com.client as win32
                        
                        # Create Word application
                        word = win32.gencache.EnsureDispatch('Word.Application')
                        word.Visible = False
                        
                        # Open document
                        doc = word.Documents.Open(tmp_file_path)
                        
                        # Extract text
                        extracted_text = doc.Content.Text
                        method_used = "win32com"
                        
                        # Close document and quit Word
                        doc.Close()
                        word.Quit()
                        
                        logger.info("Successfully extracted text using win32com")
                    except Exception as e:
                        logger.warning(f"win32com failed: {str(e)}")
                
                # Method 2: Try using textract (cross-platform)
                if not extracted_text.strip():
                    try:
                        import textract
                        extracted_text = textract.process(tmp_file_path).decode('utf-8')
                        method_used = "textract"
                        logger.info("Successfully extracted text using textract")
                    except Exception as e:
                        logger.warning(f"Textract failed: {str(e)}")
                
                # Method 3: Try reading as plain text (fallback)
                if not extracted_text.strip():
                    try:
                        # Try to read as raw text (may not work well with .doc format)
                        with open(tmp_file_path, 'rb') as f:
                            raw_content = f.read()
                        
                        # Try to decode and extract readable text
                        try:
                            # Simple text extraction from binary
                            text_content = raw_content.decode('utf-8', errors='ignore')
                            # Remove non-printable characters
                            extracted_text = ''.join(char for char in text_content if char.isprintable() or char.isspace())
                            method_used = "raw_text_extraction"
                            logger.info("Extracted text using raw text method")
                        except Exception:
                            pass
                    except Exception as e:
                        logger.warning(f"Raw text extraction failed: {str(e)}")
                
                # If all methods failed, provide helpful message
                if not extracted_text.strip():
                    return {
                        'success': False,
                        'error': 'Unable to extract text from DOC file. Please save the document as DOCX format for better compatibility, or copy/paste the content manually.',
                        'text': '',
                        'suggestion': 'Convert .doc to .docx format'
                    }
                
                # Clean up the text
                cleaned_text = extracted_text.strip()
                
                return {
                    'success': True,
                    'text': cleaned_text,
                    'method': f'doc_extraction_{method_used}',
                    'length': len(cleaned_text)
                }
                
            finally:
                # Clean up temp file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                    
        except Exception as e:
            logger.error(f"Error extracting text from DOC: {str(e)}")
            return {
                'success': False,
                'error': f'Error processing DOC file: {str(e)}. Please try converting to DOCX format.',
                'text': ''
            }
    
    @staticmethod
    def get_excel_sheet_names(file_content: bytes) -> Dict:
        """Get all sheet names from Excel file"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Read Excel file to get sheet names
                xl_file = pd.ExcelFile(tmp_file_path)
                sheet_names = xl_file.sheet_names
                
                # Close the Excel file object to release the file handle
                xl_file.close()
                
                return {
                    'success': True,
                    'sheet_names': sheet_names,
                    'total_sheets': len(sheet_names)
                }
                
            finally:
                # Clean up temp file
                try:
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                except PermissionError:
                    # If file is still locked, try again after a short delay
                    import time
                    time.sleep(0.1)
                    try:
                        if os.path.exists(tmp_file_path):
                            os.unlink(tmp_file_path)
                    except Exception as cleanup_error:
                        logger.warning(f"Could not clean up temp file {tmp_file_path}: {cleanup_error}")
                    
        except Exception as e:
            logger.error(f"Error reading Excel sheet names: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'sheet_names': []
            }
    
    @staticmethod
    def extract_data_from_excel(file_content: bytes, sheet_name: str = None, max_rows: int = 5) -> Dict:
        """Extract first N rows from Excel sheet with headers"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Read Excel file
                if sheet_name:
                    df = pd.read_excel(tmp_file_path, sheet_name=sheet_name, nrows=max_rows)
                else:
                    df = pd.read_excel(tmp_file_path, nrows=max_rows)
                
                # Convert to string representation
                data_text = df.to_string(index=False)
                
                # Also get as HTML table for better formatting
                html_table = df.to_html(index=False, classes='table table-striped')
                
                # Get basic info
                info = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist()
                }
                
                return {
                    'success': True,
                    'text': data_text,
                    'html': html_table,
                    'info': info,
                    'sheet_name': sheet_name or 'Default'
                }
                
            finally:
                # Clean up temp file with retry logic
                try:
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                except PermissionError:
                    # If file is still locked, try again after a short delay
                    import time
                    time.sleep(0.1)
                    try:
                        if os.path.exists(tmp_file_path):
                            os.unlink(tmp_file_path)
                    except Exception as cleanup_error:
                        logger.warning(f"Could not clean up temp file {tmp_file_path}: {cleanup_error}")
                    
        except Exception as e:
            logger.error(f"Error extracting data from Excel: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'text': ''
            }
    
    @staticmethod
    def extract_data_from_csv(file_content: bytes, max_rows: int = 5) -> Dict:
        """Extract first N rows from CSV file with headers"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.csv', mode='wb') as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Read CSV file
                df = pd.read_csv(tmp_file_path, nrows=max_rows)
                
                
                # Convert to string representation
                data_text = df.to_string(index=False)
                
                # Also get as HTML table
                html_table = df.to_html(index=False, classes='table table-striped')
                
                # Get basic info
                info = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist()
                }
                
                return {
                    'success': True,
                    'text': data_text,
                    'html': html_table,
                    'info': info
                }
                
            finally:
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                    
        except Exception as e:
            logger.error(f"Error extracting data from CSV: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'text': ''
            }
    
    @staticmethod
    def process_uploaded_file(filename: str, file_content: bytes) -> Dict:
        """Process uploaded file based on its extension"""
        file_ext = os.path.splitext(filename.lower())[1]
        
        if file_ext == '.docx':
            return FileProcessor.extract_text_from_docx(file_content)
        elif file_ext == '.doc':
            return FileProcessor.extract_text_from_doc(file_content)
        elif file_ext in ['.xlsx', '.xls']:
            # For Excel, first get sheet names
            sheet_info = FileProcessor.get_excel_sheet_names(file_content)
            if sheet_info['success']:
                return {
                    'success': True,
                    'type': 'excel',
                    'sheet_names': sheet_info['sheet_names'],
                    'requires_sheet_selection': len(sheet_info['sheet_names']) > 1
                }
            else:
                return sheet_info
        elif file_ext == '.csv':
            return {
                'success': True,
                'type': 'csv',
                'requires_sheet_selection': False
            }
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_ext}. Supported: .docx, .doc, .xlsx, .xls, .csv'
            }
